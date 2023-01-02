import os
import zipfile
import re
# from xml.etree.cElementTree import XML
from lxml import etree
import docx
import unicodedata
import en_ner_craft_md
import en_ner_bc5cdr_md
import en_core_web_md
from spacy import displacy
from tqdm import tqdm
import aspose.words as aw
import pandas as pd
import numpy as np

def excel_preproc(df,col):
	print(df)
	df = df.drop(columns=['S. No', "Patient's Name ",'IPD.1','UHID','Date', 'Time','Test Permformed'])[3:]
	print(df)
	df['Diagnosis']= df['Diagnosis'].str.lower()
	df = df.replace({'---': np.nan}, regex=True)
	df = df.replace('\*', '', regex=True)
	df = df.replace(',', '', regex=True)
	df['Sex'] = df['Sex'].replace('M',1)
	df['Sex'] = df['Sex'].replace('F',0)
	df['Sex'] = df['Sex'].replace('Years',0)
	df['Sex'] = df['Sex'].fillna(0)

	df['Age '] = df['Age '].replace(regex={r"\D+": 1})
	df['Age '] = df['Age '].fillna(df['Age '].mean())

	# x = df.groupby('IPD').first()
	diagnosis = df.groupby(df.index)['Diagnosis'].first()

	numeric_cols = df.columns.drop(['Diagnosis'])
	# numeric_cols = df.columns.drop(['Diagnosis','IPD'])
	df[numeric_cols] = df[numeric_cols].apply(pd.to_numeric, errors='coerce')
	x = df.groupby(df.index).agg('mean')
	# print(x)
	x[numeric_cols] = x[numeric_cols].apply(pd.to_numeric, errors='coerce')

	x = x[col]

	return x

df = pd.read_excel('Updated Rotam Records - 30-06-2021.xlsx').set_index('IPD')
patients_csv = pd.read_csv('preprocessed_patient_records.csv').set_index('IPD')
# print(patients_csv.index)
# print(patients_csv.index.dtype)
l = list(patients_csv.drop(columns=['y_ethanol','y_cld_type']).columns)
path = "files/"
dir_list = os.listdir(path)
# print(df['IPD'])
patient_id_list = []
for fname in dir_list:
	pid = int(fname[:-4][-5:])
	if pid in df.index:
		patient_id_list.append(pid)

# print(len(patient_id_list))
# print(df.loc[patient_id_list][patients_csv.drop(columns=['y_ethanol','y_cld_type']).columns])
# print(df.loc[patient_id_list])
new_patients_df = excel_preproc(df.loc[patient_id_list],l)
print(new_patients_df)

print(new_patients_df.isna().sum())
# new_patients_df = new_patients_df.fillna(x.mean())
# print(x.isna().sum())

# Finding Diagnosis and missing values for new patients from docx files
for patient_id in tqdm(patient_id_list):
	# print('Reading started doc')
	doc = aw.Document(path+patient_id+'.doc')
	# print('Reading completed doc')
	doc.save('../files_txt/'+patient_id+'.txt')
	# print('Saved txt')
	doc = doc.open('../files_docx/'+patient_id+'.txt','r')
	

patient_id_list = []

craft = en_ner_craft_md.load()
bc = en_ner_bc5cdr_md.load()
nlp = en_core_web_md.load()

# patient_id = 'IPID0073044'

for patient_id in tqdm(patient_id_list):
	# patient_id = 'IPID0073047'
	# doc = docx.Document(f'data/{patient_id} deidentified.docx')
	print('Reading started doc')
	doc = aw.Document(path+patient_id+'.doc')
	print('Reading completed doc')
	doc.save('../files_docx/'+patient_id+'.docx')
	print('Saved docx')
	doc = docx.Document('../files_docx/'+patient_id+'.docx')

	print('Found', len(doc.tables), 'tables')

	def get_table(table):
		m = len(table.row_cells(0))
		n = len(table.column_cells(0))
		f = []
		for i in range(m):
			f.append([])
			for j in range(n):
				f[i].append(table.cell(i,j).text)
		return f

	def nlp_get_terms(para):
		t = unicodedata.normalize('NFKD', para.text)
		d = {'bc5cdr':{}, 'craft':{}, 'nlp':{}}
		bt = bc(t)
		ct = craft(t)
		nt = nlp(t)
		for x in bt.ents:
			d['bc5cdr'][x.text] = x.label_
		for x in ct.ents:
			d['craft'][x.text] = x.label_
		for x in nt.ents:
			d['nlp'][x.text] = x.label_
		return d

	for para in doc.paragraphs:
		print(nlp_get_terms(para))

	# document = zipfile.ZipFile(f'data/{patient_id} deidentified.docx')
	# document_content = document.read('word/document.xml')

	# etree.parse(document)
	# .XMLParser()

	#document will be the filetype zipfile.ZipFile
	# for name in document.namelist():
	# 	print(name)
	# 	document.read(name, pwd=None)

